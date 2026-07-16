"""DOC/DOCX 文本抽取。

主路径：pywin32 调 Word COM 把 .doc 转 .docx，再用 python-docx 读取。
兜底：antiword 把 .doc 转纯文本。
.docx 直接用 python-docx 读取。
"""

from __future__ import annotations

import os
import subprocess
import tempfile
from pathlib import Path

from loguru import logger


def _doc_to_docx_via_word(doc_path: str, output_dir: str) -> str | None:
    """用 Win Word COM 把 .doc 转 .docx。失败返回 None。"""
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            abs_path = os.path.abspath(doc_path)
            doc = word.Documents.Open(abs_path)
            stem = Path(doc_path).stem
            docx_path = os.path.join(output_dir, f"{stem}.docx")
            doc.SaveAs(docx_path, FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
            doc.Close()
            return docx_path
        finally:
            word.Quit()
    except Exception as e:  # noqa: BLE001
        logger.warning(f"Word COM 转换失败：{e}")
        return None


def _doc_to_text_via_antiword(doc_path: str) -> str | None:
    """用 antiword 把 .doc 转纯文本（兜底）。失败返回 None。"""
    try:
        result = subprocess.run(
            ["antiword", str(doc_path)],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode == 0:
            return result.stdout
        logger.warning(f"antiword 失败（exit {result.returncode}）：{result.stderr[:200]}")
        return None
    except FileNotFoundError:
        logger.warning("antiword 不在 PATH 中")
        return None
    except Exception as e:  # noqa: BLE001
        logger.warning(f"antiword 异常：{e}")
        return None


def _extract_docx_text(docx_path: str) -> str:
    """用 python-docx 抽取 .docx 全文。"""
    from docx import Document  # type: ignore

    doc = Document(docx_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # 也抽取表格中的文字（试卷答案常在表格里）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    return "\n".join(paragraphs)


def extract_doc_text(doc_path: str | Path) -> str:
    """抽取 .doc 或 .docx 全文。

    .docx：直接用 python-docx。
    .doc：先用 Word COM 转 .docx，失败兜底 antiword。
    """
    doc_path = Path(doc_path)
    suffix = doc_path.suffix.lower()

    if suffix == ".docx":
        return _extract_docx_text(str(doc_path))

    # .doc：转换后读取
    with tempfile.TemporaryDirectory() as tmpdir:
        # 尝试 Word COM
        docx_path = _doc_to_docx_via_word(str(doc_path), tmpdir)
        if docx_path and os.path.exists(docx_path):
            return _extract_docx_text(docx_path)

        # 兜底 antiword
        text = _doc_to_text_via_antiword(str(doc_path))
        if text:
            return text

    raise RuntimeError(f"无法解析 .doc 文件（Word COM 和 antiword 均失败）：{doc_path}")
